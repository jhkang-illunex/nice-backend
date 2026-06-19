"""기존 보고서 docx 뒤에 '사용 질의 전문' 부록 추가 (v1, PDF 생성에 쓴 질의 그대로).

v1 _variants = 목적격 조사 '를' 통일(교정 전). 현재 PDF/DOCX 가 만들어진 그 질의.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from rag_final_report_eval import SCENARIO, SEEDS, _variants

DOCX_PATH = "docs/reports/rag/RAG_FINAL_REPORT_20260618.docx"
BLUE = RGBColor(0x2C, 0x5F, 0xA8)

_TMPL_LABEL = {
    "keyword": "키워드형", "import_q": "수입질문형", "tariff_q": "관세율형",
    "country_declare": "국가신고형", "hs_wonder": "HS궁금형", "export_sebun": "수출세번형",
}


def _set(cell, text, bold=False, size=8):
    cell.text = ""
    r = cell.paragraphs[0].add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)


def _h2(doc, text):
    r = doc.add_paragraph().add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE


def main() -> None:
    doc = Document(DOCX_PATH)
    doc.add_page_break()

    _h2(doc, "부록 A. 자연어 120문항 질의 전문")
    doc.add_paragraph("20 시드 × 6 템플릿. 본 보고서 수치는 아래 질의를 라이브 rag-server 에 그대로 호출해 산출.").runs[0].font.size = Pt(9)
    order = ["keyword", "import_q", "tariff_q", "country_declare", "hs_wonder", "export_sebun"]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["ID", "유형", "질의"]):
        _set(t.rows[0].cells[i], h, bold=True, size=9)
    for idx, (name, _pref) in enumerate(SEEDS):
        v = _variants(name, idx)
        for tmpl in order:
            row = t.add_row().cells
            _set(row[0], f"{idx:02d}-{tmpl}")
            _set(row[1], _TMPL_LABEL[tmpl])
            _set(row[2], v[tmpl])
    for c in t.rows[0].cells:
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = BLUE

    _h2(doc, "부록 B. 시나리오 13문항 질의 전문")
    cat_label = {"supply": "공급망", "unrelated": "비관련", "extra": "추가검증"}
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Light Grid Accent 1"
    for i, h in enumerate(["구분", "질의"]):
        _set(t2.rows[0].cells[i], h, bold=True, size=9)
    for cat, q in SCENARIO:
        row = t2.add_row().cells
        _set(row[0], cat_label.get(cat, cat))
        _set(row[1], q)
    for c in t2.rows[0].cells:
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = BLUE

    doc.save(DOCX_PATH)
    print(f"부록 추가 완료: {DOCX_PATH}")
    print(f"  부록A 120문항 + 부록B {len(SCENARIO)}문항")


if __name__ == "__main__":
    main()
