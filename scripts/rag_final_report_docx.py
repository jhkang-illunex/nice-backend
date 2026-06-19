"""최종 보고서 Word 생성 — /tmp/final_eval.json → docs/RAG_FINAL_REPORT_<date>.docx.

편집 가능한 네이티브 .docx (표·서식 그대로 Word 에서 수정). HTML/PDF 와 동일 내용.
"""

from __future__ import annotations

import json

from docx import Document
from docx.shared import Pt, RGBColor

DATE = "20260618"
JSON_PATH = "/tmp/final_eval.json"
DOCX_PATH = f"docs/reports/rag/RAG_FINAL_REPORT_{DATE}.docx"

V2 = {"hit@1": 0.650, "hit@5": 0.950, "hit@10": 1.000, "mrr": 0.777}
BLUE = RGBColor(0x2C, 0x5F, 0xA8)
TMPL_NAMES = {
    "keyword": "키워드형(추출 미적용)", "import_q": "수입 질문형", "tariff_q": "관세율 질문형",
    "country_declare": "국가+신고형", "hs_wonder": "HS부호 궁금형", "export_sebun": "수출+세번형",
}


def pct(x):
    return f"{x*100:.1f}%"


def set_cell(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)


def header_table(table):
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = BLUE


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE
    return p


def build(doc: Document, data: dict) -> None:
    ov, by_tmpl, by_seed = data["overall"], data["by_tmpl"], data["by_seed"]
    rows, scen = data["rows"], data["scenario"]

    title = doc.add_paragraph()
    tr = title.add_run(f"HSK RAG 자연어 질의 최종 평가 보고서 ({len(rows)}문항 + 시나리오 {len(scen)})")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = BLUE

    meta = doc.add_paragraph()
    meta.add_run("작성일 2026-06-18 · 대상 nice_rag (RRF 하이브리드 + LLM 품목추출 + agent) · 라이브 rag-server 재실행\n").font.size = Pt(9)
    meta.add_run("방법론 RAGAS(검색 랭킹: 기대 prefix rank 매칭 → hit@k·MRR). 자연어 120문항 = 20시드 × 6템플릿. 시나리오 13문항 = 공급망 6 + 비관련 5 + 추가검증 2.").font.size = Pt(9)

    # 1. 종합 지표
    h2(doc, "1. 종합 지표 (자연어 120문항, 재실행)")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = ["지표", "v2 기준 (2026-06-12)", "재실행 (2026-06-18)", "Δ"]
    for i, h in enumerate(hdr):
        set_cell(t.rows[0].cells[i], h, bold=True)
    for label, key, isp in [("hit@1", "hit@1", True), ("hit@5", "hit@5", True), ("hit@10", "hit@10", True), ("MRR", "mrr", False)]:
        ref, now = V2[key], ov[key]
        refs = pct(ref) if isp else f"{ref:.3f}"
        nows = pct(now) if isp else f"{now:.3f}"
        d = now - ref
        ds = f"{d*100:+.1f}%p" if isp else f"{d:+.3f}"
        row = t.add_row().cells
        set_cell(row[0], label)
        set_cell(row[1], refs)
        set_cell(row[2], nows, bold=True)
        set_cell(row[3], ds)
    header_table(t)
    doc.add_paragraph(f"top10 발견 {ov['found']}/{ov['n']} · latency p50 ~{data.get('latency_p50_s',0):.1f}s (문장형 LLM추출 포함, CPU) · 120문항 wall {data.get('wall120_s',0):.0f}s").runs[0].font.size = Pt(9)

    # 2. 템플릿 변형별
    h2(doc, "2. 템플릿 변형별 분해 (6종)")
    t = doc.add_table(rows=1, cols=7)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["변형", "n", "hit@1", "hit@5", "hit@10", "MRR", "발견"]):
        set_cell(t.rows[0].cells[i], h, bold=True)
    for k in ["keyword", "import_q", "tariff_q", "country_declare", "hs_wonder", "export_sebun"]:
        m = by_tmpl[k]
        row = t.add_row().cells
        for i, v in enumerate([TMPL_NAMES[k], m["n"], pct(m["hit@1"]), pct(m["hit@5"]), pct(m["hit@10"]), f"{m['mrr']:.3f}", f"{m['found']}/{m['n']}"]):
            set_cell(row[i], v)
    header_table(t)

    # 3. 시드별 MRR
    h2(doc, "3. 시드 품목별 MRR (낮은 순)")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    set_cell(t.rows[0].cells[0], "시드 품목", bold=True)
    set_cell(t.rows[0].cells[1], "MRR (6변형 평균)", bold=True)
    for name, mrr in sorted(by_seed.items(), key=lambda kv: kv[1]):
        row = t.add_row().cells
        set_cell(row[0], name)
        set_cell(row[1], f"{mrr:.3f}")
    header_table(t)

    # 4~6. 시나리오
    def scen_section(cat, title, note=""):
        items = [s for s in scen if s["cat"] == cat]
        h2(doc, f"{title} ({len(items)}문항)")
        if note:
            doc.add_paragraph(note).runs[0].font.size = Pt(9)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(["질의", "/search top3 (score)", "/agent 답변 (요약)", "판정"]):
            set_cell(t.rows[0].cells[i], h, bold=True)
        for s in items:
            top = "\n".join(f"{x['hs_code']} {x['score']} {(x['name_ko'] or '')[:14]}" for x in s["top3"]) or "(빈 리스트 — 추천 불가)"
            row = t.add_row().cells
            set_cell(row[0], s["q"], size=9)
            set_cell(row[1], top, size=8)
            set_cell(row[2], (s["answer"] or "")[:160], size=8)
            set_cell(row[3], s["verdict"], size=9)
        header_table(t)

    scen_section("supply", "4. 공급망 시나리오")
    scen_section("unrelated", "5. 비관련 일반 질의", "기대 동작: /search 빈 리스트 + /agent 명시 거부 (2중 차단).")
    scen_section("extra", "6. 추가 검증 질의", "규격 수치 보존(브릭스) · CRAG 런타임 보정(원유) 확인용.")

    # 7. 전체 테스트셋
    h2(doc, "7. 테스트셋 전체 결과 (120문항)")
    doc.add_paragraph("rank = /search top10 에서 기대 prefix 가 처음 등장한 순위 (—=top10 밖). 시드별 6변형.").runs[0].font.size = Pt(9)
    t = doc.add_table(rows=1, cols=8)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["시드", "기대 prefix", "kw", "수입", "관세", "신고", "HS", "수출"]):
        set_cell(t.rows[0].cells[i], h, bold=True)
    order = ["keyword", "import_q", "tariff_q", "country_declare", "hs_wonder", "export_sebun"]
    by_id = {r["id"]: r for r in rows}
    seeds = [(r["seed"], r.get("prefixes", [])) for r in rows if r["tmpl"] == "keyword"]
    for idx, (name, prefixes) in enumerate(seeds):
        row = t.add_row().cells
        set_cell(row[0], name, size=8)
        set_cell(row[1], ",".join(prefixes), size=8)
        for j, tm in enumerate(order):
            r = by_id.get(f"{idx:02d}-{tm}")
            rk = r["rank"] if r else None
            set_cell(row[2 + j], rk if rk is not None else "—", size=8)
    header_table(t)

    # 8. 미발견 분석
    miss = [r for r in rows if r["rank"] is None]
    refusals = [r for r in miss if not r.get("top1")]
    retr = [r for r in miss if r.get("top1")]
    h2(doc, f"8. 미발견 {len(miss)}건 분석 (거부 {len(refusals)} / 검색실패 {len(retr)})")
    doc.add_paragraph(
        "거부 = /search CRAG 가 빈 리스트(추천 불가) 반환 — 정답 점수가 CRAG 임계(0.033) 경계. "
        "검색실패 = 후보는 반환했으나 정답이 top-10 밖."
    ).runs[0].font.size = Pt(9)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["ID", "시드", "기대 prefix", "top1 (실제 1위)", "분류"]):
        set_cell(t.rows[0].cells[i], h, bold=True)
    for r in sorted(miss, key=lambda x: x["id"]):
        is_ref = not r.get("top1")
        top1 = "(빈 리스트 — 추천 불가)" if is_ref else f"{r['top1']} {(r.get('top1_name') or '')[:14]}"
        label = "거부 (CRAG 빈 리스트)" if is_ref else "검색실패 (top10 밖)"
        row = t.add_row().cells
        for i, v in enumerate([r["id"], r["seed"], ",".join(r.get("prefixes", [])), top1, label]):
            set_cell(row[i], v)
    header_table(t)
    doc.add_paragraph(
        "처방 분리 — 거부: CRAG 임계(0.033)와 정답 점수(예 LNG 271111 = 0.0328)가 겹치는 경계 문제로 "
        "임계 하향·동의어 점수 상향 시 해소. 검색실패: 색인 커버리지/임베딩 문제(예 불화수소산 281111 이 ○○수소산류보다 하위) — "
        "reranker·색인 보강(본 사업) 영역."
    ).runs[0].font.size = Pt(9)

    doc.add_paragraph(
        "방법: 라이브 rag-server /api/hsk/search(limit=10)·/api/hsk/agent(k=5) 실제 호출. "
        "hit@k=기대 prefix 가 top-k 내, MRR=1/rank. faithfulness·answer_relevancy 등 RAGAS 생성지표는 "
        "별도 judge 하니스 영역으로 본 재실행에는 검색 랭킹 지표만 포함."
    ).runs[0].font.size = Pt(8)


def main() -> None:
    with open(JSON_PATH, encoding="utf-8") as f:
        data = json.load(f)
    doc = Document()
    # 기본 폰트 — 한글
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)
    build(doc, data)
    doc.save(DOCX_PATH)
    print(f"DOCX 저장: {DOCX_PATH}")


if __name__ == "__main__":
    main()
