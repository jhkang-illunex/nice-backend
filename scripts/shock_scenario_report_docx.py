"""쇼크 시나리오 래퍼 기능 보고서 — Word(.docx) 생성 (python-docx).

PDF판(shock_scenario_report.py)과 동일 내용을 Word 로. 화면 캡처 임베드 포함.

  python scripts/shock_scenario_report_docx.py
  → docs/SHOCK_SCENARIO_REPORT_<DATE>.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DATE = "20260619"
REPORT_DIR = "docs/reports/shock"
DOCX_PATH = f"{REPORT_DIR}/SHOCK_SCENARIO_REPORT_{DATE}.docx"
IMG = {
    "overview": "/tmp/shock_flow_overview.png",
    "nodes": "/tmp/shock_zoom_nodes.png",
    "edges": "/tmp/shock_zoom_edges.png",
    "random": "/tmp/random_board.png",
    "random_delta": "/tmp/random_delta.png",
}
BLUE = RGBColor(0x2C, 0x5F, 0xA8)
IMG_W = Inches(6.3)


def _set_korean_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _mono(doc: Document, text: str) -> None:
    """계산식 블록 — 모노스페이스."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)


def _img(doc: Document, key: str) -> None:
    p = Path(IMG[key])
    if not p.exists():
        doc.add_paragraph(f"[이미지 없음: {p}]")
        return
    doc.add_picture(str(p), width=IMG_W)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _cap(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x5A, 0x6B, 0x7B)


def build() -> None:
    doc = Document()
    _set_korean_font(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.7)

    h = doc.add_heading("외생충격 시나리오 래퍼 — 기능 보고서", level=0)
    for run in h.runs:
        run.font.color.rgb = BLUE
    _cap(doc, "프로젝트: nice-backend · 모듈: nice_graph.shock · 작성일: 2026-06-19")
    _cap(doc, "대상: 관세 충격 / 거래 변화(랜덤 매출·매입 포함) + /api/shock/scenario + Streamlit 데모")

    # 0
    doc.add_heading("0. 개요 — 단일 알고리즘 + 2축 래퍼", level=1)
    doc.add_paragraph(
        "충격 전파의 실제 계산은 propagate_shock(거듭제곱급수 합) 하나뿐이다. 요구된 4갈래"
        "(관세충격×{매출,매입}, 거래변화×{매출,매입})는 새 엔진이 아니라 두 직교 축의 조합으로 "
        "구현된다. 알고리즘은 무변경."
    )
    _table(doc, ["축", "값", "의미", "구현 위치"], [
        ["방향(direction)", "upstream / downstream",
         "하류·매출 파급(매출처,A) / 상류·매입 파급(매입처,B). 엣지 방향+정규화 분모 전환", "assemble.py 인자"],
        ["시나리오(scenario)", "tariff / transaction_change",
         "W불변·시드주입 / 거래비중 g수정→변화분 Δ", "scenario.py 래퍼"],
    ])

    # 1
    doc.add_heading("1. 파이프라인 & 핵심 계산식", level=1)
    _table(doc, ["단계", "함수", "역할"], [
        ["1. 시드 선별", "select_primary_firms", "ra603 거래구성으로 HS 노출 기업 점수화"],
        ["2. 그래프 조립", "assemble_propagation_input", "company_edge depth-3 유도부분그래프 → R, init"],
        ["3. 전파", "propagate_shock", "거듭제곱급수 합(active-set 반복)"],
        ["4. 시나리오", "run_tariff_shock / run_transaction_change", "방향·가중치·g 조합 묶음"],
    ])
    doc.add_heading("1.1 전파 엔진", level=2)
    _mono(doc,
          "total_effect = Σ_{k≥0} R^k · init\n"
          "라운드 갱신:  next_shock[t] += cur_shock[s] · rate(s→t)\n"
          "종료: 모든 |propagated| ≤ ε(1e-8) → 자연수렴 / max_iter=500 안전장치")
    doc.add_heading("1.2 비중(rate) 정규화 — 방향별", level=2)
    _mono(doc,
          "α=damping, W=방향가중치(A 또는 B), amt=sly_amt(셀러→바이어)\n\n"
          "downstream (셀러s→바이어t, 매입 파급, W=B):\n"
          "   rate(s→t) = B·α·amt(s→t) / Σ_t' amt(s→t')   [분모=셀러 총매출]\n\n"
          "upstream (바이어t→셀러s, 매출 파급, W=A):\n"
          "   rate(t→s) = A·α·amt(s→t) / Σ_s' amt(s'→t)   [분모=바이어 총매입]")
    doc.add_paragraph(
        "수렴 불변식: within_subgraph 정규화로 각 source 의 Σ_out = W·α ≤ 1 → ρ(R) ≤ W·α < 1 "
        "→ 절대수렴. 방향 전환 시 분모 PARTITION 도 새 source 기준으로 전환하는 것이 핵심."
    )
    doc.add_heading("1.3 정규화 기준 옵션 (normalize) — 방향과 직교", level=2)
    doc.add_paragraph(
        "분모를 어느 끝 기준으로 잡는가를 방향과 분리해 선택. 수렴 보장과 경제적 명칭 충실 사이를 옵션화."
    )
    _table(doc, ["normalize", "분모 기준", "downstream(매입)", "upstream(매출)", "수렴"], [
        ["source(기본)", "전파 source", "셀러 총매출", "바이어 총매입", "Σ_out=W·α≤1 → 절대수렴 보장"],
        ["counterparty", "거래상대", "바이어 총매입", "셀러 총매출",
         "매출/매입 비중 라벨 충실, 단 Σ_out 무제한 → 수렴 보장 약화(발산 시 converged=False)"],
    ])
    doc.add_paragraph(
        "쌍대(dual) 불변식 검증: source 모드는 각 source Σ_out=0.85, counterparty 모드는 각 "
        "target Σ_in=0.85 가 정확히 성립(실 PG, 위반 0). rate 는 745/747 엣지에서 달라져 옵션이 "
        "실제로 계산을 바꿈을 확인."
    )

    # 2
    doc.add_heading("2. 기능 — 관세 충격 (tariff)", level=1)
    doc.add_paragraph(
        "[W 불변] 그래프 구조는 그대로, 1차 기업(시드)에 외생 충격만 주입. 한 번 호출로 "
        "매출 파급(하류·매출처, A)과 매입 파급(상류·매입처, B)을 동시 산출. "
        "(★ 문서 기준: 매출 파급=매출처(고객) 방향, 매입 파급=매입처(공급사) 방향.)"
    )
    _mono(doc,
          "init = { seed_node : shock }   (shock = score 비례 또는 균등)\n"
          "매출 파급: result_A = Σ R_downstream^k · init  (rate=A·α·매출비중)\n"
          "매입 파급: result_B = Σ R_upstream^k · init    (rate=B·α·매입비중)")
    _table(doc, ["direction", "엣지 방향", "파급 효과", "가중치", "정규화 분모"], [
        ["downstream", "셀러→바이어", "매출 파급(하류·매출처)", "A", "셀러 총매출"],
        ["upstream", "바이어→셀러", "매입 파급(상류·매입처)", "B", "바이어 총매입"],
    ])
    doc.add_heading("화면 — 시드 추출 → 시나리오 전파(방향별 탭)", level=2)
    _img(doc, "overview")
    _cap(doc, "HS 8481 → 1차 10곳 → depth-3 (노드 321·엣지 747). 매출/매입 파급 탭 생성.")
    doc.add_heading("화면 — 노드 그리드 / 엣지 그리드(rate)", level=2)
    _img(doc, "nodes")
    _cap(doc, "노드: node_id·bizno·upchecd·기업명·시드·seed_shock·shock (내림차순).")
    _img(doc, "edges")
    _cap(doc, "엣지: rate = 방향·A/B·g 가 반영된 전파 실투입값 → 래퍼 효과 직접 대조.")

    # 3
    doc.add_heading("3. 기능 — 거래 변화 (transaction_change)", level=1)
    doc.add_paragraph(
        "[W 수정] 특정 1차→2차(셀러→바이어) 거래의 비중에 0~1 인자 g 를 곱한 수정 그래프로 "
        "전파하고, 원본 대비 변화분 Δ 를 산출."
    )
    _mono(doc,
          "rate'(s→t) = g·rate(s→t)  for (s,t)∈overrides,  else rate(s→t)\n"
          "baseline = Σ R^k·init     (원 W)\n"
          "changed  = Σ R'^k·init    (수정 W)\n"
          "변화분 Δ(node) = changed(node) − baseline(node)   [difference-of-runs]")
    doc.add_paragraph(
        "실측 예 (override 1438123482→3148200884, g=0.5, downstream): 변화 노드 130곳 중 "
        "128곳 감소(Δ<0), 대상 바이어 Δ=−0.0199 — 부호·방향이 경제 직관과 일치."
    )

    doc.add_heading("3.1 1차↔2차 매출/매입 랜덤 가중치 (신규)", level=2)
    doc.add_paragraph(
        "특정 HS 에 연계된 1차 기업이 2차 기업과 맺은 거래의 매출/매입에 랜덤 g 를 자동 부여. "
        "엣지를 일일이 지정하지 않고 한 번에 생성하며, 난수 시드로 재현 가능."
    )
    _table(doc, ["분류", "거래 방향(저장: 셀러→바이어)", "대상 엣지"], [
        ["매출(sales)", "1차(셀러) → 2차(바이어)", "sb∈1차, bb∉1차"],
        ["매입(purchase)", "2차(셀러) → 1차(바이어)", "bb∈1차, sb∉1차"],
    ])
    _mono(doc,
          "후보 = { (s,t)∈1차↔2차 : side=sales→매출만 / purchase→매입만 / both→둘 다 }\n"
          "정렬(후보) 후 g(s,t)=Uniform(low,high) with seed  (DB 행순서 무관·재현 보장)\n"
          "1차↔1차·2차↔3차 (양끝 동시 1차 또는 1차 미포함)는 제외")
    doc.add_paragraph(
        "옵션: side(both/sales/purchase) · 범위[low,high]⊆[0,1](상한 1→수렴 유지) · "
        "seed(재현) · only_firms(일부 1차 한정, 비우면 전체)."
    )
    doc.add_paragraph(
        "API: POST /api/shock/scenario 에 random_override:{side,low,high,seed,only_firms} 추가. "
        "지정 시 서버가 1차↔2차 매출/매입 랜덤 g 생성, 응답 applied_overrides 로 실제 적용값 반환."
    )
    doc.add_heading("화면 — 거래 변화 보드(랜덤 모드)", level=2)
    _img(doc, "random")
    _cap(doc, "대상거래·g범위·재현시드·대상1차 → 랜덤 가중치 생성 → 구분(매출/매입)·셀러·바이어·g 그리드. "
              "실측 HS 8481·seed 42 → 40건(매출 24·매입 16).")
    doc.add_heading("화면 — 변화분 Δ 결과(랜덤 적용 후)", level=2)
    _img(doc, "random_delta")
    _cap(doc, "방향별 탭·변화분Δ 그리드(음수=거래축소로 인한 파급 감소)·CSV. 실측 Σ변화분Δ=−15.898.")

    # 4
    doc.add_heading("4. 엔드포인트 — POST /api/shock/scenario", level=1)
    _table(doc, ["필드", "타입", "설명"], [
        ["scenario", "tariff | transaction_change", "시나리오 종류"],
        ["seeds", "[{bizno, upchecd, shock}]", "1차 기업 + 초기충격"],
        ["directions", "[upstream, downstream]", "계산 방향(기본 둘 다)"],
        ["weight_a / weight_b", "float (>0)", "매출(하류·A) / 매입(상류·B) 가중치"],
        ["normalize", "source | counterparty", "분모 기준 — source(수렴보장)/counterparty(매출·매입 비중)"],
        ["edge_overrides", "[{from_bizno, to_bizno, factor}]", "거래변화 — 명시 g"],
        ["random_override", "{side, low, high, seed, only_firms}", "거래변화 — 1차↔2차 매출/매입 랜덤 g"],
    ])
    _mono(doc,
          "응답:\n"
          "{ scenario, warnings[], applied_overrides:[{from_bizno,to_bizno,factor}],\n"
          "  directions:[{ direction, effect_label, weight,\n"
          "               shock_list:[{bizno,shock}], total_shock, iterations, converged,\n"
          "               n_nodes, n_edges }] }")

    # 5
    doc.add_heading("5. 기능 검증 결과", level=1)
    doc.add_paragraph(
        "자동화 테스트 69개 전부 통과(propagate 18 + 라우터 15 + 시나리오 36). "
        "추가로 실제 PG 기반 수학 불변식·HTTP 응답 교차검증. "
        "코드리뷰 반영(시드 이중계상·입력검증·중복 assemble 제거·열거/디스패치 통합) 후 회귀 0."
    )
    doc.add_heading("5.1 수학 불변식 (실 PG · HS 8481 · 노드 321 · 엣지 747)", level=2)
    _table(doc, ["불변식", "결과", "측정"], [
        ["downstream 각 source Σ_out == 0.85", "PASS", "위반 0 · 정규화 분모 정확"],
        ["upstream 각 source Σ_out == 0.85", "PASS", "위반 0 · 방향전환 후 재정규화 정확"],
        ["노드 집합 down == up", "PASS", "차집합 0 (321=321)"],
        ["엣지 방향 뒤집힘(down == reversed up)", "PASS", "불일치 0 (747=747)"],
        ["direction_weight=0.5 → rate 절반", "PASS", "A/B 선형 반영"],
        ["override g=0.4 → 대상 엣지만 ×0.4", "PASS", "0.0159→0.0064"],
        ["Δ == 수동(changed − baseline)", "PASS", "difference-of-runs 정확"],
        ["거래비중 축소 → 하류 Δ<0", "PASS", "대상 바이어 Δ=−0.0199"],
        ["tariff 양방향 수렴", "PASS", "up 104회/Σ32.08 · down 103회/Σ43.01"],
        ["랜덤 override 모두 1차↔2차", "PASS", "매출 24·매입 16·위반 0"],
        ["랜덤 seed 재현성 / side·only_firms 필터", "PASS", "동일 seed→동일 40건 · 필터 정확"],
        ["랜덤 거래변화 전파 수렴·감소우세", "PASS", "128/128 Δ<0 · Σ=−19.07"],
        ["normalize=source: source Σ_out=0.85", "PASS", "위반 0"],
        ["normalize=counterparty: target Σ_in=0.85(dual)", "PASS", "위반 0 · rate 745/747 변경"],
    ])
    doc.add_heading("5.2 HTTP 엔드포인트 (실 DB)", level=2)
    _table(doc, ["검증", "결과"], [
        ["tariff 200 + 양방향(A=0.8/B=0.6) 구조·라벨·수렴", "PASS"],
        ["transaction_change 200 + Δ(130중 128 감소, Σ=−0.67)", "PASS"],
        ["transaction_change 빈 overrides → 422", "PASS"],
        ["random_override 경로 → 생성·전달·applied_overrides 직렬화", "PASS"],
        ["OpenAPI 경로 등록", "PASS"],
    ])
    doc.add_paragraph(
        "종합: 기능 정상 동작. 렌더링뿐 아니라 방향 전환·정규화·가중치·거래변화 Δ·랜덤 생성의 "
        "수치 정확성까지 확인됨."
    )

    # 6
    doc.add_heading("6. 도메인 정의 — 확정 사항", level=1)
    doc.add_paragraph(
        "방향↔라벨(2026-06-19 화면기획안 기준 확정): 매출 파급=매출처(고객)/하류/downstream/A, "
        "매입 파급=매입처(공급사)/상류/upstream/B. 정규화도 정합(매출=셀러 총매출, 매입=바이어 총매입). "
        "엔진(orientation·normalize) 무변경, 라벨·가중치 바인딩만 문서에 맞춤.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "변화분 정의(확정): 거래변화=거래내역(매입/매출 비중)에 g 반영(W→W′), 변화분=그 순효과="
        "difference-of-runs(수정W′−원W)의 Δ. 명세 '변화분을 seed로'='결과로 나오는 값이 곧 변화분'.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "전파 거래연도: 1차 시드 선택 후 company_edge.trade_year(전체/2024/2026) 선택 "
        "— screen 의 기준연도(ra603 bse_yr)와 별개 축.",
        style="List Bullet",
    )

    Path(REPORT_DIR).mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"DOCX 작성: {DOCX_PATH}")


if __name__ == "__main__":
    build()
